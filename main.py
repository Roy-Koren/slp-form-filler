from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import requests
from pathlib import Path

app = FastAPI()

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration
DEEPGRAM_API_KEY = os.getenv("DEEPGRAM_API_KEY", "")
os.makedirs("documents", exist_ok=True)

# Form questions
QUESTIONS = [
    {"id": "patient_name", "label": "שם המטופל", "question": "מה שם המטופל?"},
    {"id": "treatment_date", "label": "תאריך הטיפול", "question": "מה היום של הטיפול?"},
    {"id": "session_number", "label": "מס' מפגש", "question": "זה מפגש מספר כמה?"},
    {"id": "goal_1", "label": "מטרה 1", "question": "מה המטרה הראשונה?"},
    {"id": "goal_2", "label": "מטרה 2", "question": "מה המטרה השנייה?"},
    {"id": "goal_3", "label": "מטרה 3", "question": "מה המטרה השלישית?"},
    {"id": "session_description", "label": "תיאור הטיפול", "question": "תאר מה קרה בטיפול (עד 5 משפטים)"},
    {"id": "functioning_goal1", "label": "תפקוד - מטרה 1", "question": "איך התנהל המטופל ביחס למטרה הראשונה?"},
    {"id": "functioning_goal2", "label": "תפקוד - מטרה 2", "question": "איך התנהל המטופל ביחס למטרה השנייה?"},
    {"id": "functioning_goal3", "label": "תפקוד - מטרה 3", "question": "איך התנהל המטופל ביחס למטרה השלישית?"},
    {"id": "learnings", "label": "מה למדתי", "question": "מה למדת על המטופל? (עד 4 שורות)"},
    {"id": "therapist_reflection", "label": "דמות טיפולית", "question": "אני כדמות טיפולית - מה הרגשת, מה עשית נכון, מה יכולת לעשות אחרת? (עד 4 שורות)"},
    {"id": "supervisor_question", "label": "שאלה למדריכה", "question": "יש לך שאלה למדריכה?"},
]

@app.get("/", response_class=HTMLResponse)
async def get_home():
    """Serve the web interface"""
    html_path = Path(__file__).parent / "index.html"
    if html_path.exists():
        with open(html_path, "r", encoding="utf-8") as f:
            return f.read()
    return "<h1>SLP Form Filler</h1>"

@app.post("/api/transcribe")
async def transcribe(audio: UploadFile = File(...), question_id: str = Form(...)):
    """
    Transcribe audio from browser using Deepgram
    """
    try:
        if not DEEPGRAM_API_KEY:
            return {
                "question_id": question_id,
                "transcription": "[Set DEEPGRAM_API_KEY environment variable]",
                "status": "no_api_key"
            }
        
        # Read audio file
        audio_content = await audio.read()
        
        print(f"📝 Transcribing audio for {question_id} - {len(audio_content)} bytes")
        
        # Send to Deepgram with correct headers
        headers = {
            "Authorization": f"Token {DEEPGRAM_API_KEY}",
            "Content-Type": "audio/wav"
        }
        
        response = requests.post(
            "https://api.deepgram.com/v1/listen?model=nova-3&language=he",
            headers=headers,
            data=audio_content,
            timeout=30
        )
        
        print(f"Deepgram response status: {response.status_code}")
        
        if response.status_code != 200:
            print(f"❌ Deepgram error: {response.text}")
            return {
                "question_id": question_id,
                "transcription": f"[Error {response.status_code}]",
                "status": "error"
            }
        
        result = response.json()
        
        # Extract transcript
        try:
            transcript = result["results"]["channels"][0]["alternatives"][0]["transcript"]
            print(f"✅ Transcribed: {transcript}")
        except (KeyError, IndexError):
            print(f"❌ Could not extract transcript from: {result}")
            transcript = "[No speech detected]"
        
        return {
            "question_id": question_id,
            "transcription": transcript,
            "status": "success"
        }
    
    except Exception as e:
        print(f"❌ Transcription error: {e}")
        return {
            "question_id": question_id,
            "transcription": f"[Error: {str(e)[:30]}]",
            "status": "error"
        }

@app.post("/api/generate-document")
async def generate_document(form_data: dict):
    """Generate Word document from form answers using original template"""
    try:
        template_path = Path(__file__).parent / "template_דוח_יומי.docx"
        
        # Load original template
        if template_path.exists():
            doc = Document(str(template_path))
        else:
            # Fallback if template not found
            doc = Document()
        
        # Fill in the form fields in paragraphs
        for para in doc.paragraphs:
            text = para.text
            
            # Replace field placeholders
            for question in QUESTIONS:
                field_id = question["id"]
                answer = form_data.get(field_id, "")
                
                if answer:
                    # Replace in paragraph text
                    if field_id in text or question["label"] in text:
                        # This handles simple text replacement
                        para.text = para.text.replace(question["label"] + ":", 
                                                     f"{question['label']}: {answer}")
        
        # Also fill in table cells if they exist
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for question in QUESTIONS:
                            field_id = question["id"]
                            answer = form_data.get(field_id, "")
                            
                            if answer and field_id in para.text:
                                para.text = answer
        
        # Save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_name = f"דוח_יומי_{timestamp}.docx"
        file_path = os.path.join("documents", file_name)
        doc.save(file_path)
        
        print(f"✅ Document saved: {file_path}")
        
        return {
            "status": "success",
            "file_name": file_name,
            "file_path": file_path
        }
    
    except Exception as e:
        print(f"❌ Document error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{file_name}")
async def download_document(file_name: str):
    """Download generated document"""
    file_path = os.path.join("documents", file_name)
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=file_name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.get("/api/questions")
async def get_questions():
    """Get all form questions"""
    return QUESTIONS

@app.get("/health")
async def health():
    """Health check"""
    return {"status": "ok"}

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    print(f"🚀 Starting SLP Form Filler on port {port}")
    print(f"📱 Open http://localhost:{port} in your browser")
    if not DEEPGRAM_API_KEY:
        print(f"⚠️  WARNING: DEEPGRAM_API_KEY not set!")
    uvicorn.run(app, host="0.0.0.0", port=port)
