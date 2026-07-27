from docx import Document
from pathlib import Path
from datetime import datetime
import os

# Sample form data
test_form_data = {
    "patient_name": "דוד כהן",
    "treatment_date": "27 יולי 2024",
    "session_number": "5",
    "goal_1": "שיפור הגייה של קול ר",
    "goal_2": "הגברת מודעות לשפל נשימה",
    "goal_3": "תרגול בליעה בטוחה",
    "session_description": "עבדנו על הגייה עם תרגילים קוליים. הבחור הראה עניין גדול והשתתף בפעילויות.",
    "functioning_goal1": "דוד הצליח בתרגיל הראשון",
    "functioning_goal2": "הראה מודעות טובה",
    "functioning_goal3": "בליעה שיפרה",
    "learnings": "דוד מתגובב טוב לטיפול ומאוד שיתופי",
    "therapist_reflection": "הרגשתי שהקשר טוב, עשיתי אותו מרגיע, יכולתי להשתמש בחזותיים יותר",
    "supervisor_question": "איך אני יכול לשפר את האמצעים?",
    "next_goal_1": "המשך עבודה על הגייה",
    "next_method_1": "באמצעות תרגילים ביתיים",
    "next_goal_2": "עבודה על קול",
    "next_method_2": "תרגילי קול מובנים",
    "next_goal_3": "עבודה על בליעה",
    "next_method_3": "בליעה בתנוחות שונות",
}

# Load template
template_path = Path("template_דוח_יומי.docx")

if not template_path.exists():
    print("Template not found!")
    exit(1)

doc = Document(str(template_path))

# Fill header paragraphs
for para in doc.paragraphs:
    text = para.text
    
    if "שם המטופל:" in text and "patient_name" in test_form_data:
        para.text = "שם המטופל: " + test_form_data['patient_name'] + "\t\t\t נוכחים:\t\t המטפל:"
    
    if "תאריך הטיפול:" in text and "treatment_date" in test_form_data:
        para.text = "תאריך הטיפול: " + test_form_data['treatment_date'] + "\t\t מס' מפגש: " + test_form_data.get('session_number', '')

# Fill table cells
if len(doc.tables) > 0:
    table = doc.tables[0]
    
    # Goals (rows 0-2)
    for i in range(3):
        goal_key = "goal_" + str(i+1)
        if goal_key in test_form_data and len(table.rows) > i:
            table.rows[i].cells[0].text = test_form_data[goal_key]
    
    # Session description (row 3)
    if "session_description" in test_form_data and len(table.rows) > 3:
        table.rows[3].cells[0].text = test_form_data['session_description']
    
    # Patient functioning (rows 4-6)
    for i in range(3):
        func_key = "functioning_goal" + str(i+1)
        if func_key in test_form_data and len(table.rows) > 4+i:
            table.rows[4+i].cells[0].text = test_form_data[func_key]
    
    # Learnings (row 7)
    if "learnings" in test_form_data and len(table.rows) > 7:
        table.rows[7].cells[0].text = test_form_data['learnings']
    
    # Therapist reflection (row 8)
    if "therapist_reflection" in test_form_data and len(table.rows) > 8:
        table.rows[8].cells[0].text = test_form_data['therapist_reflection']
    
    # Supervisor question (row 9)
    if "supervisor_question" in test_form_data and len(table.rows) > 9:
        table.rows[9].cells[0].text = test_form_data['supervisor_question']
    
    # Next goals and methods (rows 11-13)
    for i in range(3):
        goal_key = "next_goal_" + str(i+1)
        method_key = "next_method_" + str(i+1)
        
        if len(table.rows) > 11+i:
            if goal_key in test_form_data:
                table.rows[11+i].cells[1].text = test_form_data[goal_key]
            if method_key in test_form_data:
                table.rows[11+i].cells[0].text = test_form_data[method_key]

# Save
os.makedirs("documents", exist_ok=True)
timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
file_name = "דוח_יומי_TEST_" + timestamp + ".docx"
file_path = os.path.join("documents", file_name)
doc.save(file_path)

print("Test document created: " + file_path)
